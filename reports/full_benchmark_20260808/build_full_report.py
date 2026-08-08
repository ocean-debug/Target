"""Build reproducible metrics, figures, and a Word report for the full local benchmark run."""
from __future__ import annotations

import json
import math
import platform
import statistics
import sys
import xml.etree.ElementTree as ET
from collections import Counter, defaultdict
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUT = Path(__file__).resolve().parent
RAW = OUT / "raw"
FIGURES = OUT / "figures"
DOCX = OUT / "Target_full_benchmark_report_20260808.docx"

sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "src"))

from benchmark.evaluate_context_relations import evaluate  # noqa: E402

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "DCEAF7"
LIGHT_GRAY = "F2F4F7"
DARK = "172033"
MUTED = "5B6573"
GREEN = "2E8B57"
ORANGE = "D77A27"


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def read_jsonl(path: Path):
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def percentile(values: list[float], p: float) -> float:
    values = sorted(values)
    if not values:
        return 0.0
    index = (len(values) - 1) * p
    low, high = math.floor(index), math.ceil(index)
    if low == high:
        return values[low]
    return values[low] + (values[high] - values[low]) * (index - low)


def parse_pytest(path: Path) -> dict:
    root = ET.parse(path).getroot()
    suite = root if root.tag == "testsuite" else root.find("testsuite")
    tests = int(suite.attrib.get("tests", 0))
    failures = int(suite.attrib.get("failures", 0))
    errors = int(suite.attrib.get("errors", 0))
    skipped = int(suite.attrib.get("skipped", 0))
    return {
        "tests": tests,
        "passed": tests - failures - errors - skipped,
        "skipped": skipped,
        "failures": failures,
        "errors": errors,
        "time_s": round(float(suite.attrib.get("time", 0.0)), 3),
    }


def relation_predictions(gold: list[dict]) -> tuple[list[dict], list[dict]]:
    blind = []
    oracle = []
    for entry in gold:
        blind_label = "supported_anchor" if entry["task_family"] == "disease_target_anchor" else "context_complete"
        blind.append({
            "id": entry["id"],
            "label": blind_label,
            "actions": ["preserve_context", "trace_evidence"],
            "claims": [],
        })
        oracle.append({
            "id": entry["id"],
            "label": entry["gold"]["label"],
            "actions": entry["gold"]["required_actions"],
            "claims": [],
        })
    return blind, oracle


def write_jsonl(path: Path, rows: list[dict]) -> None:
    path.write_text("".join(json.dumps(row, ensure_ascii=False, sort_keys=True) + "\n" for row in rows), encoding="utf-8")


def collect_metrics() -> dict:
    main = load_json(RAW / "main" / "benchmark_report.json")
    diseases = load_json(RAW / "diseases" / "benchmark_report.json")
    pytest = parse_pytest(RAW / "pytest.xml")
    relation_gold = read_jsonl(ROOT / "benchmark" / "goldset_context_relations.jsonl")
    blind_rows, oracle_rows = relation_predictions(relation_gold)
    write_jsonl(RAW / "relation_context_blind_predictions.jsonl", blind_rows)
    write_jsonl(RAW / "relation_oracle_predictions.jsonl", oracle_rows)
    blind_report = evaluate(relation_gold, blind_rows)
    oracle_report = evaluate(relation_gold, oracle_rows)
    (RAW / "relation_context_blind_report.json").write_text(
        json.dumps(blind_report, ensure_ascii=False, indent=2, sort_keys=True) + "\n", encoding="utf-8"
    )
    (RAW / "relation_oracle_report.json").write_text(
        json.dumps(oracle_report, ensure_ascii=False, indent=2, sort_keys=True) + "\n", encoding="utf-8"
    )

    disease_times = [float(task.get("elapsed_s", 0)) for task in diseases["tasks"] if task.get("passed") is not None]
    bucket_times: dict[str, list[float]] = defaultdict(list)
    for task in diseases["tasks"]:
        if task.get("passed") is not None:
            bucket_times[task["category"]].append(float(task.get("elapsed_s", 0)))

    metrics = {
        "run_date": "2026-08-08",
        "environment": {
            "python": platform.python_version(),
            "platform": platform.platform(),
            "note": "Local Windows run; project acceptance runtime is Python 3.11.",
        },
        "pytest": pytest,
        "main_benchmark": main,
        "disease_benchmark": diseases,
        "disease_runtime": {
            "mean_s": round(statistics.mean(disease_times), 3),
            "median_s": round(statistics.median(disease_times), 3),
            "p95_s": round(percentile(disease_times, 0.95), 3),
            "max_s": round(max(disease_times), 3),
            "by_bucket": {
                key: {
                    "mean_s": round(statistics.mean(vals), 3),
                    "median_s": round(statistics.median(vals), 3),
                    "p95_s": round(percentile(vals, 0.95), 3),
                }
                for key, vals in sorted(bucket_times.items())
            },
        },
        "relation_benchmark": {
            "cases": len(relation_gold),
            "families": dict(Counter(row["task_family"] for row in relation_gold)),
            "splits": dict(Counter(row["split"] for row in relation_gold)),
            "labels": dict(Counter(row["gold"]["label"] for row in relation_gold)),
            "context_blind_baseline": blind_report,
            "oracle_integrity_check": oracle_report,
        },
        "scope_boundary": {
            "live_tasks_run": 0,
            "reason": "Live external-API tasks and Reviewer LoRA matrix require network/API/GPU deployment profiles and are informational, not CI gates.",
        },
    }
    (OUT / "metrics_summary.json").write_text(
        json.dumps(metrics, ensure_ascii=False, indent=2, sort_keys=True) + "\n", encoding="utf-8"
    )
    return metrics


def setup_plot_font() -> None:
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            font_manager.fontManager.addfont(str(candidate))
            plt.rcParams["font.family"] = font_manager.FontProperties(fname=str(candidate)).get_name()
            break
    plt.rcParams["axes.unicode_minus"] = False


def savefig(name: str) -> Path:
    path = FIGURES / name
    plt.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def create_figures(metrics: dict) -> dict[str, Path]:
    FIGURES.mkdir(parents=True, exist_ok=True)
    setup_plot_font()
    figs: dict[str, Path] = {}

    labels = ["pytest 用例", "主 Benchmark 任务", "疾病库任务", "关系 Gold 自检"]
    passed = [
        metrics["pytest"]["passed"],
        metrics["main_benchmark"]["summary"]["tasks_passed"],
        metrics["disease_benchmark"]["summary"]["tasks_passed"],
        metrics["relation_benchmark"]["cases"],
    ]
    totals = [
        metrics["pytest"]["tests"] - metrics["pytest"]["skipped"],
        metrics["main_benchmark"]["summary"]["tasks"],
        metrics["disease_benchmark"]["summary"]["tasks"],
        metrics["relation_benchmark"]["cases"],
    ]
    rates = [100 * p / t for p, t in zip(passed, totals)]
    fig, ax = plt.subplots(figsize=(9, 4.8))
    bars = ax.barh(labels, rates, color=["#2E74B5", "#3D8DFF", "#6DCBF4", "#2E8B57"])
    ax.set_xlim(0, 105)
    ax.set_xlabel("通过率（%）")
    ax.set_title("完整本地验收：四个阶段均通过")
    ax.grid(axis="x", color="#E6EAF0")
    ax.set_axisbelow(True)
    for bar, p, t in zip(bars, passed, totals):
        ax.text(101, bar.get_y() + bar.get_height() / 2, f"{p}/{t}", va="center", ha="right", fontweight="bold")
    figs["overview"] = savefig("01_validation_overview.png")

    main_categories = metrics["main_benchmark"]["summary"]["categories"]
    disease_categories = metrics["disease_benchmark"]["summary"]["categories"]
    cat_labels = ["主链", "迁移", "确定性/恢复", "鲁棒性", "合同", "疾病正常", "缺失上下文", "冲突证据", "因果陷阱"]
    category_rows = list(main_categories.values()) + list(disease_categories.values())
    scores = [100 * row["score"] for row in category_rows]
    assertions = [row["assertions"] for row in category_rows]
    fig, ax = plt.subplots(figsize=(10, 5.2))
    bars = ax.bar(cat_labels, scores, color=["#2E74B5"] * 5 + ["#6DCBF4"] * 4)
    ax.set_ylim(0, 112)
    ax.set_ylabel("断言通过率（%）")
    ax.set_title("关键能力维度：261 条断言全部通过")
    ax.tick_params(axis="x", rotation=28)
    ax.grid(axis="y", color="#E6EAF0")
    ax.set_axisbelow(True)
    for bar, count in zip(bars, assertions):
        ax.text(bar.get_x() + bar.get_width() / 2, 102, f"n={count}", ha="center", va="bottom", fontsize=9)
    figs["categories"] = savefig("02_category_assertions.png")

    blind = metrics["relation_benchmark"]["context_blind_baseline"]["metrics"]
    oracle = metrics["relation_benchmark"]["oracle_integrity_check"]["metrics"]
    metric_keys = ["coverage", "label", "required_actions", "forbidden_claims"]
    metric_labels = ["覆盖率", "标签准确率", "必要动作召回", "禁用声明安全"]
    x = list(range(len(metric_keys)))
    width = 0.36
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.bar([i - width / 2 for i in x], [100 * blind[k] for k in metric_keys], width, label="上下文盲基线", color="#D77A27")
    ax.bar([i + width / 2 for i in x], [100 * oracle[k] for k in metric_keys], width, label="Gold/评分器自检", color="#2E8B57")
    ax.set_xticks(x, metric_labels)
    ax.set_ylim(0, 112)
    ax.set_ylabel("得分（%）")
    ax.set_title("关系评测能区分“识别锚点”与“理解上下文”")
    ax.legend(frameon=False)
    ax.grid(axis="y", color="#E6EAF0")
    ax.set_axisbelow(True)
    for i, key in enumerate(metric_keys):
        ax.text(i - width / 2, 100 * blind[key] + 2, f"{100 * blind[key]:.1f}", ha="center", fontsize=9)
        ax.text(i + width / 2, 100 * oracle[key] + 2, f"{100 * oracle[key]:.0f}", ha="center", fontsize=9)
    figs["relation_metrics"] = savefig("03_relation_baseline_vs_oracle.png")

    split_counts = metrics["relation_benchmark"]["splits"]
    label_counts = metrics["relation_benchmark"]["labels"]
    fig, axes = plt.subplots(1, 2, figsize=(10, 4.6))
    axes[0].bar(split_counts.keys(), split_counts.values(), color=["#2E74B5", "#6DCBF4", "#A9DDF4"])
    axes[0].set_title("疾病不跨 split")
    axes[0].set_ylabel("样本数")
    axes[1].bar(["锚点", "完整", "缺失", "错配"], [
        label_counts["supported_anchor"], label_counts["context_complete"],
        label_counts["insufficient_context"], label_counts["context_mismatch"],
    ], color=["#2E74B5", "#2E8B57", "#D7A52A", "#D77A27"])
    axes[1].set_title("Gold 标签构成")
    for ax in axes:
        ax.grid(axis="y", color="#E6EAF0")
        ax.set_axisbelow(True)
        for bar in ax.patches:
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 1, str(int(bar.get_height())), ha="center")
    figs["relation_composition"] = savefig("04_relation_composition.png")

    main_tasks = [task for task in metrics["main_benchmark"]["tasks"] if task.get("passed") is not None]
    fig, ax = plt.subplots(figsize=(9, 5.2))
    names = [task["id"] for task in main_tasks]
    times = [float(task.get("elapsed_s", 0)) for task in main_tasks]
    bars = ax.barh(names, times, color="#3D8DFF")
    ax.invert_yaxis()
    ax.set_xlabel("耗时（秒）")
    ax.set_title("主 Benchmark 单任务耗时：确定性三连跑最慢")
    ax.grid(axis="x", color="#E6EAF0")
    ax.set_axisbelow(True)
    for bar, value in zip(bars, times):
        ax.text(value + 0.05, bar.get_y() + bar.get_height() / 2, f"{value:.2f}", va="center", fontsize=9)
    figs["main_runtime"] = savefig("05_main_task_runtime.png")

    bucket_times: dict[str, list[float]] = defaultdict(list)
    for task in metrics["disease_benchmark"]["tasks"]:
        if task.get("passed") is not None:
            bucket_times[task["category"]].append(float(task.get("elapsed_s", 0)))
    ordered = [
        "disease_library_normal", "disease_library_missing_context",
        "disease_library_conflicting_evidence", "disease_library_trap",
    ]
    fig, ax = plt.subplots(figsize=(9, 4.8))
    ax.boxplot([bucket_times[k] for k in ordered], tick_labels=["正常", "缺失上下文", "冲突证据", "因果陷阱"], patch_artist=True,
               boxprops={"facecolor": "#DCEAF7", "edgecolor": "#2E74B5"}, medianprops={"color": "#D77A27", "linewidth": 2})
    ax.set_ylabel("单任务耗时（秒）")
    ax.set_title("72 个疾病任务耗时分布稳定")
    ax.grid(axis="y", color="#E6EAF0")
    ax.set_axisbelow(True)
    figs["disease_runtime"] = savefig("06_disease_runtime_distribution.png")

    return figs


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError(f"table widths must total 9360 DXA, got {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_run_font(run, size=None, bold=None, color=None, name="Microsoft YaHei") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text: str, bold=False, color=DARK, after=6, size=11, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    for run in p.runs:
        set_run_font(run, size=11, color=DARK)
    if not p.runs:
        run = p.add_run(text)
        set_run_font(run, size=11, color=DARK)
    else:
        p.runs[0].text = text
    return p


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, LIGHT_GRAY)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=10, bold=True, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            for p in cells[index].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=9.5, color=DARK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc, path: Path, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.35))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = True
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("TARGET · FULL BENCHMARK REPORT")
    set_run_font(run, size=9, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("2026-08-08 · Local reproducible run")
    set_run_font(run, size=9, color=MUTED)


def create_docx(metrics: dict, figs: dict[str, Path]) -> None:
    doc = Document()
    configure_document(doc)

    kicker = add_text(doc, "TARGET DISCOVERY AGENT", bold=True, color=BLUE, after=12, size=10)
    kicker.paragraph_format.space_before = Pt(8)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(8)
    title_p.paragraph_format.keep_with_next = True
    run = title_p.add_run("完整评测运行与关系评测报告")
    set_run_font(run, size=26, bold=True, color=DARK)
    add_text(doc, "主链、疾病库、上下文关系与回归测试的统一验收", color=MUTED, after=14, size=13)

    add_table(doc, ["项目", "值"], [
        ["运行日期", metrics["run_date"]],
        ["分支", "agent/context-relation-benchmark"],
        ["提交基线", "939d9b9"],
        ["运行模式", "本地 fake/unit；未执行 live/GPU"],
        ["Python", metrics["environment"]["python"]],
    ], [2700, 6660])

    metric_table = doc.add_table(rows=2, cols=4)
    metric_table.style = "Table Grid"
    values = [
        ("11/11", "主 Benchmark"),
        ("72/72", "疾病任务"),
        (f"{metrics['pytest']['passed']}/{metrics['pytest']['tests'] - metrics['pytest']['skipped']}", "pytest 通过"),
        ("145", "关系样本"),
    ]
    for i, (value, label) in enumerate(values):
        metric_table.cell(0, i).text = value
        metric_table.cell(1, i).text = label
        set_cell_shading(metric_table.cell(0, i), LIGHT_BLUE)
        for run in metric_table.cell(0, i).paragraphs[0].runs:
            set_run_font(run, size=18, bold=True, color=DARK_BLUE)
        for run in metric_table.cell(1, i).paragraphs[0].runs:
            set_run_font(run, size=9, bold=True, color=MUTED)
        metric_table.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        metric_table.cell(1, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(metric_table, [2340, 2340, 2340, 2340])

    add_heading(doc, "执行结论", 1)
    add_text(doc, "本次完整本地可复现验收全部通过，没有发现 fake/unit 主链、疾病矩阵或既有单元测试回归。新增关系评测能明显区分“只识别疾病靶点”与“理解组织、细胞类型和疾病时期上下文”两种能力。", after=8)
    add_bullet(doc, "主 Benchmark：11/11 任务、27/27 断言通过。")
    add_bullet(doc, "疾病库：72/72 任务、234/234 断言通过，覆盖正常、缺失上下文、冲突证据与因果陷阱。")
    add_bullet(doc, f"pytest：{metrics['pytest']['passed']} 通过、{metrics['pytest']['skipped']} 跳过、0 失败。")
    add_bullet(doc, "关系评测：上下文盲基线的标签准确率和必要动作召回均为 62.8%；Gold/评分器自检为 100%。")
    add_figure(doc, figs["overview"], "图 1｜完整本地验收阶段通过率。关系 Gold 为评分器完整性检查，不代表 Agent 实际性能。")
    add_heading(doc, "判读原则", 2)
    add_bullet(doc, "主 Benchmark、疾病库与 pytest 是当前 Agent fake/unit 执行结果，可用于本地回归门禁。")
    add_bullet(doc, "关系上下文盲基线是刻意设计的参照系；62.8% 表明数据集能识别缺失上下文与上下文错配。")
    add_bullet(doc, "Gold/评分器自检只证明样本与 scorer 对齐，不应表述为 Agent 达到 100% 的关系理解能力。")

    doc.add_page_break()
    add_heading(doc, "1. 评测范围与运行方法", 1)
    add_heading(doc, "1.1 本次执行范围", 2)
    add_table(doc, ["阶段", "规模", "判定对象", "结果"], [
        ["pytest", f"{metrics['pytest']['tests']} 用例", "代码与合同回归", f"{metrics['pytest']['passed']} 通过 / {metrics['pytest']['skipped']} 跳过"],
        ["主 Benchmark", "11 任务 / 27 断言", "主链、迁移、确定性、鲁棒性、合同", "100%"],
        ["疾病库", "72 任务 / 234 断言", "18 疾病 × 4 任务桶", "100%"],
        ["关系评测", "145 样本", "疾病—靶点—组织—细胞—时期", "基线 62.8% / 自检 100%"],
    ], [1800, 1900, 3700, 1960])
    add_heading(doc, "1.2 未执行范围", 2)
    add_text(doc, "Live 外部 API 任务和 Reviewer LoRA live matrix 未在本地运行。它们需要网络/API/GPU 部署配置，且仓库 rubric 将 live 结果定义为信息性结果而非 CI 合并门槛。")
    add_text(doc, "运行环境提示：本机 Python 版本为 " + metrics["environment"]["python"] + "；仓库正式 acceptance runtime 为 Python 3.11，因此发布前仍应在远端 3.11 profile 复跑。", bold=True, color=ORANGE)

    add_heading(doc, "2. 主 Benchmark 指标", 1)
    add_figure(doc, figs["categories"], "图 2｜主 Benchmark 与疾病库各能力维度断言通过率；柱顶 n 为断言数。")
    main_summary = metrics["main_benchmark"]["summary"]
    add_text(doc, f"主 Benchmark 共执行 {main_summary['tasks']} 个 fake/unit 任务，{main_summary['assertions']} 条断言全部通过。它验证了 UC 主链输出、Legacy/LangGraph 一致性、三连跑确定性、终态与中断恢复、预算降级、MCH 边界、OOD 拒绝以及合同/Schema 门槛。")
    add_figure(doc, figs["main_runtime"], "图 3｜主 Benchmark 单任务耗时。BM-03 包含三次确定性运行，因此耗时最高。")

    doc.add_page_break()
    add_heading(doc, "3. 疾病库 Benchmark 指标", 1)
    disease_summary = metrics["disease_benchmark"]["summary"]
    add_text(doc, f"18 个疾病分别进入 normal、missing_context、conflicting_evidence 与 trap 四类任务，共 {disease_summary['tasks']} 项、{disease_summary['assertions']} 条断言，全部通过。")
    add_table(doc, ["任务桶", "任务数", "断言", "通过率"], [
        ["normal", "18", "54", "100%"],
        ["missing_context", "18", "54", "100%"],
        ["conflicting_evidence", "18", "54", "100%"],
        ["trap", "18", "72", "100%"],
    ], [3100, 1700, 1900, 2660])
    add_figure(doc, figs["disease_runtime"], "图 4｜四类疾病任务的单任务耗时分布。")
    runtime = metrics["disease_runtime"]
    add_text(doc, f"72 个疾病任务平均耗时 {runtime['mean_s']:.3f} 秒，中位数 {runtime['median_s']:.3f} 秒，P95 为 {runtime['p95_s']:.3f} 秒，最大值 {runtime['max_s']:.3f} 秒。首个 UC 上下文因初始化成本略高，其余任务分布集中。")

    add_heading(doc, "4. 上下文关系评测", 1)
    add_text(doc, "关系评测不把“组织/时期置换”解释成生物学关系为假，而只判断输入是否偏离当前 curated benchmark context。这样可以测出上下文保留、缺失追问、错配降级和因果边界，同时避免制造不可靠的跨疾病 hard negatives。")
    add_figure(doc, figs["relation_composition"], "图 5｜145 条关系样本的 split 与 Gold 标签构成。所有同一疾病样本只进入一个 split。")
    add_figure(doc, figs["relation_metrics"], "图 6｜上下文盲基线与 Gold/评分器自检。100% 安全得分来自基线没有输出禁用声明，不代表其上下文判断正确。")
    add_table(doc, ["指标", "上下文盲基线", "Gold/评分器自检", "解释"], [
        ["覆盖率", "100.0%", "100.0%", "两组均输出全部 145 条"],
        ["标签准确率", "62.8%", "100.0%", "盲基线无法识别缺失与错配"],
        ["必要动作召回", "62.8%", "100.0%", "盲基线不会触发追问/降级"],
        ["禁用声明安全", "100.0%", "100.0%", "盲基线不输出禁止声明"],
    ], [2300, 1900, 2000, 3160])

    add_heading(doc, "5. 关键步骤指标总览", 1)
    add_table(doc, ["步骤", "核心指标", "结果", "判定"], [
        ["Goldset 生成", "字节一致、Schema valid", "145/145", "通过"],
        ["任务规划与主链", "BM-01 主链断言", "11/11", "通过"],
        ["迁移兼容", "Legacy/LangGraph parity", "3/3", "通过"],
        ["确定性与恢复", "deterministic/resume", "3/3", "通过"],
        ["鲁棒性", "预算降级 + OOD", "4/4", "通过"],
        ["合同与 Schema", "version/whitelist/export", "3/3", "通过"],
        ["疾病泛化", "18 疾病 × 4 桶", "234/234", "通过"],
        ["关系上下文", "盲基线 label/action", "62.8% / 62.8%", "具有区分度"],
        ["关系评分器", "Gold 自检四指标", "100%", "通过"],
        ["代码回归", "pytest", f"{metrics['pytest']['passed']} pass / {metrics['pytest']['skipped']} skip", "通过"],
    ], [2100, 3300, 2100, 1860])

    add_heading(doc, "6. 结论与下一步", 1)
    add_text(doc, "当前提交已经形成可合并的本地质量闭环：原有主链和疾病矩阵保持全绿，新增关系评测提供了明确的上下文敏感度信号。它适合立即用于 Planner/Reviewer 输出适配器的回归测试，但尚不能替代论文级人工 Gold 或真实 live 工具评估。")
    add_bullet(doc, "接入 Agent 输出适配器，报告真实 Planner 的 relation label、required actions 与 forbidden claims。")
    add_bullet(doc, "为每个 disease-target anchor 增加 PMID/DOI、方向证据、cell/tissue/stage source span。")
    add_bullet(doc, "落实 paper-level 与 2021–2024 / 2025 / 2026 time split，避免论文派生资产泄漏。")
    add_bullet(doc, "在远端 Python 3.11 acceptance profile 与 Reviewer LoRA GPU profile 复跑。")

    add_heading(doc, "附录 A｜复现命令", 1)
    for command in [
        "python -m pytest -q --junitxml=reports/full_benchmark_20260808/raw/pytest.xml",
        "python benchmark/runner.py --goldset benchmark/goldset_v2.jsonl --out reports/full_benchmark_20260808/raw/main",
        "python benchmark/generate_disease_goldset.py --check",
        "python benchmark/runner.py --goldset benchmark/goldset_diseases.jsonl --out reports/full_benchmark_20260808/raw/diseases",
        "python benchmark/generate_context_relation_goldset.py --check",
        "python reports/full_benchmark_20260808/build_full_report.py",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(command)
        set_run_font(run, name="Consolas", size=8.5, color=DARK)

    add_heading(doc, "附录 B｜数据来源", 1)
    for source in [
        "benchmark/goldset_v2.jsonl 与 benchmark/rubric.md",
        "benchmark/goldset_diseases.jsonl 与 configs/disease_library.yaml",
        "benchmark/goldset_context_relations.jsonl 与 schemas/context_relation_case.schema.json",
        "reports/full_benchmark_20260808/raw/*（本次执行原始报告与 JUnit XML）",
    ]:
        add_bullet(doc, source)

    doc.core_properties.title = "Target 完整评测运行与关系评测报告"
    doc.core_properties.subject = "Full local benchmark, disease matrix, context relation benchmark"
    doc.core_properties.author = "Target team"
    doc.core_properties.comments = "Generated from committed benchmark outputs; no external claims added."
    doc.save(DOCX)


def main() -> int:
    metrics = collect_metrics()
    figs = create_figures(metrics)
    create_docx(metrics, figs)
    print(DOCX)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
