"""Build whole-project metrics, figures, and an expanded Word report."""
from __future__ import annotations

import json
import math
import re
import statistics
import sys
from collections import Counter, defaultdict
from copy import deepcopy
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import yaml
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
OUT = Path(__file__).resolve().parent
FIGURES = OUT / "figures"
BASE_DOCX = OUT / "Target_full_benchmark_report_20260808.docx"
DOCX = OUT / "Target_project_full_visualization_report_20260808.docx"
METRICS_JSON = OUT / "project_metrics.json"

sys.path.insert(0, str(OUT))
import build_full_report as base  # noqa: E402


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def load_jsonl(path: Path):
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def collect_code_metrics() -> dict:
    top_lines = Counter()
    extension_lines = Counter()
    python_files = []
    test_functions = {}
    excluded = {".git", ".pytest_cache", "reports", "runs_archive"}
    for path in ROOT.rglob("*"):
        if not path.is_file() or any(part in excluded for part in path.parts):
            continue
        try:
            lines = path.read_text(encoding="utf-8").splitlines()
        except (UnicodeDecodeError, OSError):
            continue
        nonblank = sum(bool(line.strip()) for line in lines)
        rel = path.relative_to(ROOT)
        top_lines[rel.parts[0]] += nonblank
        extension_lines[path.suffix or "<none>"] += nonblank
        if path.suffix == ".py":
            python_files.append({"path": rel.as_posix(), "nonblank_lines": nonblank})
        if rel.parts[0] == "tests" and path.name.startswith("test_"):
            test_functions[path.name] = len(re.findall(r"^\s*def test_", path.read_text(encoding="utf-8"), re.M))
    return {
        "nonblank_lines_by_area": dict(top_lines.most_common()),
        "nonblank_lines_by_extension": dict(extension_lines.most_common()),
        "largest_python_files": sorted(python_files, key=lambda row: -row["nonblank_lines"])[:15],
        "test_functions_by_file": dict(sorted(test_functions.items())),
        "test_functions": sum(test_functions.values()),
    }


def collect_project_metrics() -> dict:
    disease_doc = yaml.safe_load((ROOT / "configs/disease_library.yaml").read_text(encoding="utf-8"))
    diseases = disease_doc["diseases"]
    tools = yaml.safe_load((ROOT / "configs/tool_registry.yaml").read_text(encoding="utf-8"))["tools"]
    workflow = yaml.safe_load((ROOT / "workflows/disease_to_target.yaml").read_text(encoding="utf-8"))

    schemas = []
    for path in sorted((ROOT / "schemas").glob("*.json")):
        payload = load_json(path)
        schemas.append({
            "name": path.stem.replace(".schema", ""),
            "properties": len(payload.get("properties", {})),
            "required": len(payload.get("required", [])),
        })

    alignment_files = {
        "SFT": ROOT / "alignment_data/reviewer_sft.jsonl",
        "Preference": ROOT / "alignment_data/reviewer_preferences.jsonl",
        "Heldout": ROOT / "alignment_data/acceptance_heldout.jsonl",
    }
    alignment = {}
    for name, path in alignment_files.items():
        rows = load_jsonl(path)
        alignment[name] = {
            "rows": len(rows),
            "categories": dict(Counter(row["category"] for row in rows)),
            "risk": dict(Counter(row["risk"] for row in rows)),
            "splits": dict(Counter(row["split"] for row in rows)),
        }

    entries = sorted((ROOT / "runs_archive/matrix_full/entries").iterdir())
    status = Counter()
    artifact_presence = Counter()
    tool_calls, target_cards, rank_lengths, evidence_counts = [], [], [], []
    decisions, ranked_genes = Counter(), Counter()
    score_dimensions = defaultdict(list)
    finding_categories, finding_severity, finding_resolved = Counter(), Counter(), Counter()
    for entry in entries:
        for filename in ("status.json", "ranked_targets.json", "reviewer_findings.jsonl", "report.md"):
            artifact_presence[filename] += int((entry / filename).exists())
        run_status = load_json(entry / "status.json")
        status[run_status["terminal_status"]] += 1
        tool_calls.append(run_status["detail"].get("tool_calls", 0))
        target_cards.append(run_status["detail"].get("target_cards", 0))
        ranked = load_json(entry / "ranked_targets.json")
        rank_lengths.append(len(ranked))
        for target in ranked:
            decisions[target.get("decision", "UNKNOWN")] += 1
            ranked_genes[target["gene"]] += 1
            evidence_counts.append(len(target.get("evidence_ids", [])))
            for key, value in target.get("scores", {}).items():
                if key not in {"contract_version", "total"} and isinstance(value, (int, float)):
                    score_dimensions[key].append(float(value))
        for finding in load_jsonl(entry / "reviewer_findings.jsonl"):
            finding_categories[finding["category"]] += 1
            finding_severity[finding["severity"]] += 1
            finding_resolved[str(bool(finding["resolved"]))] += 1

    uc_candidates = load_json(ROOT / "data/derived/uc_candidates_v2.json")
    perturbation = load_json(ROOT / "data/derived/uc_observed_perturbation_v2.json")
    mch = load_json(ROOT / "data/derived/mch_gold_v2.json")
    candidate_genes = {row["gene"] for row in uc_candidates["candidates"]}
    perturbation_genes = {row["gene"] for row in perturbation["targets"]}

    metrics = {
        "inventory": {
            "diseases": len(diseases),
            "reference_targets": sum(len(d["reference_targets"]) for d in diseases),
            "tools": len(tools),
            "enabled_tools": sum(bool(t["enabled"]) for t in tools),
            "schemas": len(schemas),
            "workflow_states": len(workflow["states"]),
            "required_tools": len(workflow["required_tools"]),
            "alignment_rows": sum(v["rows"] for v in alignment.values()),
            "archived_runs": len(entries),
        },
        "workflow": {
            "states": workflow["states"],
            "terminal_states": workflow["terminal_states"],
            "limits": workflow["limits"],
            "required_tools": workflow["required_tools"],
        },
        "tools": {
            "enabled": [t["id"] for t in tools if t["enabled"]],
            "disabled": [t["id"] for t in tools if not t["enabled"]],
        },
        "disease_library": {
            "categories": dict(Counter(d["category"] for d in diseases)),
            "tissues": dict(Counter(d["context"]["tissue"] for d in diseases)),
            "cell_types": dict(Counter(d["context"]["cell_type"] for d in diseases)),
            "reference_evidence": dict(Counter(t["evidence"] for d in diseases for t in d["reference_targets"])),
            "reference_genes": dict(Counter(t["gene"] for d in diseases for t in d["reference_targets"])),
        },
        "schemas": schemas,
        "alignment": alignment,
        "code": collect_code_metrics(),
        "derived_data": {
            "uc_candidates": uc_candidates,
            "perturbation": perturbation,
            "candidate_perturbation_overlap": len(candidate_genes & perturbation_genes),
            "mch": mch,
        },
        "archive": {
            "terminal_status": dict(status),
            "artifact_presence": dict(artifact_presence),
            "tool_calls": {"mean": statistics.mean(tool_calls), "min": min(tool_calls), "max": max(tool_calls)},
            "target_cards": dict(Counter(target_cards)),
            "ranked_targets_per_run": dict(Counter(rank_lengths)),
            "decisions": dict(decisions),
            "top_ranked_genes": dict(ranked_genes.most_common(20)),
            "evidence_ids_per_target": {
                "mean": round(statistics.mean(evidence_counts), 3),
                "median": statistics.median(evidence_counts),
                "max": max(evidence_counts),
            },
            "score_dimensions": {
                key: {"mean": round(statistics.mean(values), 3), "max": round(max(values), 3)}
                for key, values in score_dimensions.items()
            },
            "reviewer_findings": {
                "categories": dict(finding_categories),
                "severity": dict(finding_severity),
                "resolved": dict(finding_resolved),
                "total": sum(finding_categories.values()),
            },
        },
        "scope_note": "Archive matrix is a stored fake/unit replay. All 72 runs ended completed_with_gaps; it is not a live biological validation cohort.",
    }
    METRICS_JSON.write_text(json.dumps(metrics, ensure_ascii=False, indent=2, sort_keys=True) + "\n", encoding="utf-8")
    return metrics


def savefig(name: str) -> Path:
    path = FIGURES / name
    plt.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def annotate_bars(ax):
    for bar in ax.patches:
        value = bar.get_height() if bar.get_height() >= 0 else 0
        ax.text(bar.get_x() + bar.get_width() / 2, value + max(0.3, ax.get_ylim()[1] * 0.012),
                f"{value:.0f}", ha="center", va="bottom", fontsize=8)


def create_project_figures(m: dict) -> dict[str, Path]:
    base.setup_plot_font()
    FIGURES.mkdir(exist_ok=True, parents=True)
    figs = {}

    inv = m["inventory"]
    labels = ["疾病", "工具", "Schema", "工作流状态", "训练/验收样本", "归档运行"]
    values = [inv["diseases"], inv["tools"], inv["schemas"], inv["workflow_states"], inv["alignment_rows"], inv["archived_runs"]]
    fig, ax = plt.subplots(figsize=(10, 4.8))
    bars = ax.bar(labels, values, color=["#2E74B5", "#3D8DFF", "#6DCBF4", "#2E8B57", "#D7A52A", "#D77A27"])
    ax.set_title("Target 项目数据与契约资产全景")
    ax.set_ylabel("数量")
    ax.grid(axis="y", color="#E6EAF0"); ax.set_axisbelow(True)
    for bar, value in zip(bars, values):
        ax.text(bar.get_x()+bar.get_width()/2, value+3, str(value), ha="center", fontweight="bold")
    figs["inventory"] = savefig("07_project_inventory.png")

    disease = m["disease_library"]
    fig, axes = plt.subplots(1, 3, figsize=(12, 4.3))
    for ax, (title, counts) in zip(axes, [
        ("疾病类别", disease["categories"]), ("组织覆盖", disease["tissues"]), ("细胞类型覆盖", disease["cell_types"])
    ]):
        rows = sorted(counts.items(), key=lambda kv: (-kv[1], kv[0]))[:8]
        ax.barh([x[0] for x in rows][::-1], [x[1] for x in rows][::-1], color="#3D8DFF")
        ax.set_title(title); ax.grid(axis="x", color="#E6EAF0"); ax.set_axisbelow(True)
    figs["disease_landscape"] = savefig("08_disease_context_landscape.png")

    fig, axes = plt.subplots(1, 2, figsize=(11, 4.5))
    evidence = disease["reference_evidence"]
    axes[0].bar(evidence.keys(), evidence.values(), color="#2E8B57")
    axes[0].set_title("73 个参考靶点的证据等级"); axes[0].tick_params(axis="x", rotation=25); axes[0].grid(axis="y", color="#E6EAF0")
    genes = sorted(disease["reference_genes"].items(), key=lambda kv: (-kv[1], kv[0]))[:12]
    axes[1].barh([x[0] for x in genes][::-1], [x[1] for x in genes][::-1], color="#2E74B5")
    axes[1].set_title("跨疾病重复出现的参考基因"); axes[1].grid(axis="x", color="#E6EAF0")
    for ax in axes: ax.set_axisbelow(True)
    figs["reference_evidence"] = savefig("09_reference_evidence_and_genes.png")

    fig, axes = plt.subplots(1, 2, figsize=(11, 4.7))
    areas = list(m["code"]["nonblank_lines_by_area"].items())[:8]
    axes[0].barh([x[0] for x in areas][::-1], [x[1] for x in areas][::-1], color="#2E74B5")
    axes[0].set_title("代码/契约非空行分布"); axes[0].grid(axis="x", color="#E6EAF0")
    schema_rows = sorted(m["schemas"], key=lambda row: -row["properties"])[:10]
    axes[1].barh([r["name"] for r in schema_rows][::-1], [r["properties"] for r in schema_rows][::-1], color="#6DCBF4")
    axes[1].set_title("Schema 属性数量（复杂度代理）"); axes[1].grid(axis="x", color="#E6EAF0")
    for ax in axes: ax.set_axisbelow(True)
    figs["code_contracts"] = savefig("10_code_and_contract_footprint.png")

    cats = list(next(iter(m["alignment"].values()))["categories"].keys())
    fig, ax = plt.subplots(figsize=(10, 4.7))
    bottom = [0]*len(cats)
    colors = ["#2E74B5", "#6DCBF4", "#D7A52A"]
    for color, (name, info) in zip(colors, m["alignment"].items()):
        vals = [info["categories"].get(cat, 0) for cat in cats]
        ax.bar(cats, vals, bottom=bottom, label=name, color=color)
        bottom = [a+b for a,b in zip(bottom, vals)]
    ax.set_title("Reviewer 对齐数据：六类高风险场景完全均衡")
    ax.set_ylabel("样本数"); ax.tick_params(axis="x", rotation=22); ax.legend(frameon=False)
    ax.grid(axis="y", color="#E6EAF0"); ax.set_axisbelow(True)
    figs["alignment"] = savefig("11_alignment_data_composition.png")

    candidates = m["derived_data"]["uc_candidates"]["candidates"]
    fig, ax = plt.subplots(figsize=(9, 5.2))
    cell_types = sorted({row["cell_type"] for row in candidates})
    palette = {cell: color for cell, color in zip(cell_types, ["#2E74B5", "#D77A27", "#2E8B57", "#7A5AA6", "#6DCBF4"])}
    for cell in cell_types:
        rows = [r for r in candidates if r["cell_type"] == cell]
        ax.scatter([r["log2fc"] for r in rows], [-math.log10(max(r["fdr"], 1e-300)) for r in rows],
                   s=[30+r["disease_strength_0_60"]*4 for r in rows], alpha=.78, label=cell, color=palette[cell])
    for row in sorted(candidates, key=lambda r: r["disease_strength_0_60"], reverse=True)[:6]:
        ax.annotate(row["gene"], (row["log2fc"], -math.log10(max(row["fdr"],1e-300))), fontsize=8, xytext=(4,4), textcoords="offset points")
    ax.axvline(0, color="#9AA4B2", lw=1); ax.set_xlabel("log2 fold-change"); ax.set_ylabel("-log10(FDR)")
    ax.set_title("UC 20 候选基因：效应方向、显著性与疾病强度")
    ax.legend(frameon=False, fontsize=8); ax.grid(color="#E6EAF0"); ax.set_axisbelow(True)
    figs["uc_candidates"] = savefig("12_uc_candidate_omics.png")

    perturb = m["derived_data"]["perturbation"]["targets"]
    fig, ax = plt.subplots(figsize=(9, 5.2))
    ax.scatter([r["activation_log2fc"] for r in perturb], [r["disease_alignment"] for r in perturb],
               s=[18+min(r["n_cells"],2000)/25 for r in perturb], color="#3D8DFF", alpha=.58)
    candidate_set = {r["gene"] for r in candidates}
    highlight = [r for r in perturb if r["gene"] in candidate_set]
    ax.scatter([r["activation_log2fc"] for r in highlight], [r["disease_alignment"] for r in highlight],
               s=65, color="#D77A27", label="UC candidate overlap")
    for row in sorted(highlight, key=lambda r: abs(r["disease_alignment"]), reverse=True)[:8]:
        ax.annotate(row["gene"], (row["activation_log2fc"], row["disease_alignment"]), fontsize=8, xytext=(4,4), textcoords="offset points")
    ax.axhline(0, color="#9AA4B2", lw=1); ax.set_xlabel("CRISPRa activation log2FC"); ax.set_ylabel("disease alignment")
    ax.set_title("71 个扰动靶点：激活强度与疾病方向一致性")
    ax.legend(frameon=False); ax.grid(color="#E6EAF0"); ax.set_axisbelow(True)
    figs["perturbation"] = savefig("13_uc_perturbation_landscape.png")

    archive = m["archive"]
    fig, ax = plt.subplots(figsize=(10, 4.8))
    stages = ["status", "ranking", "findings", "report", "target cards"]
    vals = [archive["artifact_presence"]["status.json"], archive["artifact_presence"]["ranked_targets.json"],
            archive["artifact_presence"]["reviewer_findings.jsonl"], archive["artifact_presence"]["report.md"],
            sum(int(k)*v for k,v in archive["target_cards"].items())/5]
    bars = ax.bar(stages, vals, color=["#2E74B5", "#3D8DFF", "#D77A27", "#2E8B57", "#6DCBF4"])
    ax.set_ylim(0,80); ax.set_title("72 个归档运行的端到端产物完整性"); ax.set_ylabel("完整运行数")
    ax.grid(axis="y", color="#E6EAF0"); ax.set_axisbelow(True)
    for bar,val in zip(bars,vals): ax.text(bar.get_x()+bar.get_width()/2,val+1,str(int(val)),ha="center",fontweight="bold")
    figs["archive_pipeline"] = savefig("14_archive_pipeline_completeness.png")

    fig, axes = plt.subplots(1, 2, figsize=(11, 4.7))
    dims = archive["score_dimensions"]
    axes[0].bar(dims.keys(), [v["mean"] for v in dims.values()], color=["#2E74B5", "#6DCBF4", "#7A5AA6", "#2E8B57", "#D77A27", "#D7A52A"])
    axes[0].set_title("720 个排名结果的平均维度分"); axes[0].tick_params(axis="x", rotation=30); axes[0].grid(axis="y", color="#E6EAF0")
    decisions = archive["decisions"]
    axes[1].bar(decisions.keys(), decisions.values(), color=["#2E8B57", "#D7A52A", "#D77A27"])
    axes[1].set_title("排名决策分布"); axes[1].tick_params(axis="x", rotation=18); axes[1].grid(axis="y", color="#E6EAF0")
    for ax in axes: ax.set_axisbelow(True)
    figs["ranking"] = savefig("15_ranking_dimensions_and_decisions.png")

    fig, axes = plt.subplots(1, 2, figsize=(11, 4.7))
    findings = archive["reviewer_findings"]
    axes[0].bar(findings["categories"].keys(), findings["categories"].values(), color="#D77A27")
    axes[0].set_title(f"Reviewer findings：共 {findings['total']} 条"); axes[0].tick_params(axis="x", rotation=25); axes[0].grid(axis="y", color="#E6EAF0")
    axes[1].bar(findings["severity"].keys(), findings["severity"].values(), color=["#D7A52A", "#D77A27"])
    axes[1].set_title("严重程度"); axes[1].grid(axis="y", color="#E6EAF0")
    for ax in axes: ax.set_axisbelow(True)
    figs["reviewer"] = savefig("16_reviewer_findings.png")

    top = list(archive["top_ranked_genes"].items())[:15]
    fig, ax = plt.subplots(figsize=(9, 5.4))
    ax.barh([x[0] for x in top][::-1], [x[1] for x in top][::-1], color="#2E74B5")
    ax.set_title("归档矩阵中最常进入 Top-10 的基因（非生物学验证）"); ax.set_xlabel("出现次数")
    ax.grid(axis="x", color="#E6EAF0"); ax.set_axisbelow(True)
    figs["top_genes"] = savefig("17_archive_top_ranked_genes.png")

    mch = m["derived_data"]["mch"]
    fig, ax = plt.subplots(figsize=(7.5, 4.7))
    values = [100*mch["paper"]["direction_prediction"]["accuracy"], 100*mch["project_replication"]["direction_prediction"]["accuracy"]]
    bars = ax.bar(["Nature paper", "project replication"], values, color=["#2E8B57", "#2E74B5"])
    ax.set_ylim(0,100); ax.set_ylabel("direction accuracy (%)"); ax.set_title("MCH 因果金标准：论文与项目复现")
    ax.grid(axis="y", color="#E6EAF0"); ax.set_axisbelow(True)
    for bar,val in zip(bars,values): ax.text(bar.get_x()+bar.get_width()/2,val+2,f"{val:.1f}%",ha="center",fontweight="bold")
    figs["mch"] = savefig("18_mch_replication.png")
    return figs


def update_cover(doc: Document, m: dict) -> None:
    doc.paragraphs[1].runs[0].text = "Target 项目全景与完整评测报告"
    doc.paragraphs[2].runs[0].text = "架构、数据资产、Agent 工作流、运行产物、训练对齐与 Benchmark 的统一可视化"
    values = [
        (str(m["inventory"]["diseases"]), "疾病"),
        (str(m["inventory"]["tools"]), "工具"),
        (str(m["inventory"]["schemas"]), "Schema"),
        (str(m["inventory"]["archived_runs"]), "归档运行"),
    ]
    table = doc.tables[1]
    for i, (value, label) in enumerate(values):
        table.cell(0, i).text = value; table.cell(1, i).text = label
        for run in table.cell(0,i).paragraphs[0].runs: base.set_run_font(run,size=18,bold=True,color=base.DARK_BLUE)
        for run in table.cell(1,i).paragraphs[0].runs: base.set_run_font(run,size=9,bold=True,color=base.MUTED)
        table.cell(0,i).paragraphs[0].alignment = base.WD_ALIGN_PARAGRAPH.CENTER
        table.cell(1,i).paragraphs[0].alignment = base.WD_ALIGN_PARAGRAPH.CENTER
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if "FULL BENCHMARK REPORT" in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace("FULL BENCHMARK REPORT", "PROJECT PANORAMA REPORT")


def build_docx(m: dict, figs: dict[str, Path]) -> None:
    doc = Document(BASE_DOCX)
    update_cover(doc, m)
    appendix = next(p for p in doc.paragraphs if p.text.startswith("附录 A"))
    body = doc._element.body
    marker = doc.add_paragraph("PROJECT_SECTION_MARKER")
    marker_el = marker._p
    marker.add_run().add_break(WD_BREAK.PAGE)

    base.add_heading(doc, "7. 项目全景与 Agent 架构", 1)
    base.add_text(doc, "Target 的主体不是评测平台，而是一个可追溯的疾病驱动靶点发现 Agent。项目把 TaskSpec、Planner、工具路由、证据存储、Reviewer、六维排序、TargetCards 与报告串成一条受契约约束的工作流。")
    base.add_figure(doc, figs["inventory"], "图 7｜项目核心数据、契约、训练与归档运行资产数量。")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    base.add_table(doc, ["层", "规模", "职责"], [
        ["Workflow", f"{m['inventory']['workflow_states']} 状态", "Intake 到 Report 的显式状态机"],
        ["Tools", f"{m['inventory']['enabled_tools']}/{m['inventory']['tools']} 启用", "只允许 registry 中的工具；3 个兼容插件默认关闭"],
        ["Contracts", f"{m['inventory']['schemas']} Schema", "约束任务、工具、证据、Reviewer、排序和报告"],
        ["Scientific gates", "4 类 claim", "FACT / OBSERVED / PREDICTED / INFERRED 分离"],
    ], [1700, 1900, 5760])
    base.add_text(doc, "工作流上限：30 次工具调用、20 个初始候选、2 轮 Review；上下文匹配分低于 0.5 的结果不得进入正式排名。", bold=True, color=base.ORANGE)

    base.add_heading(doc, "8. 疾病、组织、细胞与参考靶点资产", 1)
    base.add_figure(doc, figs["disease_landscape"], "图 8｜18 个疾病覆盖 5 类疾病、13 种组织与 14 种细胞类型。")
    base.add_figure(doc, figs["reference_evidence"], "图 9｜73 个参考靶点的证据等级与跨疾病重复基因。参考靶点用于 sanity check，不代表新发现。")
    base.add_text(doc, "参考证据以 approved_drug 为主（42/73），其次为 GWAS（14）、Mendelian（10）、clinical trial（6）与 mechanistic（1）。这使疾病库适合做可解释回归，但也意味着其锚点偏向已有知识。")

    base.add_heading(doc, "9. 代码、工具与契约结构", 1)
    base.add_figure(doc, figs["code_contracts"], "图 10｜项目代码/契约足迹与 Schema 复杂度。src、benchmark 与 schemas 是主体。")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    base.add_table(doc, ["对象", "数量", "观察"], [
        ["Python 非空行", f"{m['code']['nonblank_lines_by_extension'].get('.py',0):,}", "omics.py 是最大的实现文件"],
        ["测试函数", str(m['code']['test_functions']), "覆盖 runtime、contracts、tools、LoRA 与 benchmark"],
        ["工具注册", f"{m['inventory']['enabled_tools']} 启用 / {len(m['tools']['disabled'])} 关闭", "关闭项均为显式兼容插件"],
        ["Schema", str(m['inventory']['schemas']), "ToolResult、DatasetCandidate 与 EvidenceItem 属性最多"],
    ], [2100, 1900, 5360])

    base.add_heading(doc, "10. 派生生物数据与因果金标准", 1)
    base.add_figure(doc, figs["uc_candidates"], "图 11｜UC 20 候选基因的观测组学分布。点大小表示 disease strength；差异表达不是因果证据。")
    base.add_figure(doc, figs["perturbation"], "图 12｜71 个 CRISPRa 扰动靶点的激活效应与疾病方向一致性；橙色为 UC 候选交集。")
    base.add_figure(doc, figs["mch"], "图 13｜MCH 因果配置中，论文方向预测准确率为 72.9%，项目复现为 63.9%（94/147；置换 p=0.00019998）。")
    base.add_text(doc, "UC 数据来自公开快照 GSE125527 与 GSE190604；MCH/K562 是项目唯一明确的 causal gold 配置。所有数据均带 limitations，不能把关联或模型预测包装成实验事实。", bold=True, color=base.ORANGE)

    base.add_heading(doc, "11. 归档 Agent 运行：流程产物、Reviewer 与排序", 1)
    base.add_figure(doc, figs["archive_pipeline"], "图 14｜72 个归档 fake/unit 运行均生成 status、Top-10 排名、Reviewer findings、5 张 TargetCards 与报告。")
    archive = m["archive"]
    base.add_table(doc, ["重要步骤", "每运行指标", "全矩阵结果"], [
        ["ToolExecution", f"{archive['tool_calls']['mean']:.0f} 次工具调用", "72/72 一致"],
        ["Evidence / Ranking", f"10 个排名；平均 {archive['evidence_ids_per_target']['mean']:.2f} evidence IDs/target", "720 个排名结果"],
        ["Reviewer", f"平均 {archive['reviewer_findings']['total']/72:.1f} findings", f"共 {archive['reviewer_findings']['total']} 条"],
        ["TargetCard / Report", "5 张卡片 + 1 份报告", "72/72 产物完整"],
        ["Terminal", "completed_with_gaps", "72/72"],
    ], [2300, 3500, 3560])
    base.add_figure(doc, figs["reviewer"], "图 15｜Reviewer findings 以 dataset_ineligibility、context_mismatch 与 coverage_gap 为主；归档中 1,403 条均未标记 resolved。")
    base.add_figure(doc, figs["ranking"], "图 16｜归档 Top-10 的六维平均分与决策分布。disease omics、perturbation、safety_translation 在该 fake/unit 矩阵中均为 0。")
    base.add_figure(doc, figs["top_genes"], "图 17｜归档矩阵中最常进入 Top-10 的基因。频次来自固定回放资产，不是跨疾病生物学验证。")
    base.add_text(doc, "项目级结论比 benchmark 的“全绿”更严格：执行和产物合同完整，但 72 个归档运行全部是 completed_with_gaps；组学、扰动与安全维度在该矩阵中没有形成得分，Reviewer 缺口也未闭环。这些是下一阶段最重要的工程与科学信号。", bold=True, color=base.ORANGE)

    base.add_heading(doc, "12. Reviewer 对齐与验收数据", 1)
    base.add_figure(doc, figs["alignment"], "图 18｜210 条高风险 Reviewer 数据按六类场景均衡分布：120 SFT、60 preference、30 heldout。")
    base.add_text(doc, "六类场景分别是 missing_context、out_of_distribution、conflicting_evidence、causal_overreach、tool_failure 与 correct_refusal。manifest 明确禁止自动训练；所有高风险样本在训练或晋级前需要生命科学与工程双重审批。")

    base.add_heading(doc, "13. 项目级结论与优先级", 1)
    for text in [
        "产品闭环已存在：工作流、工具、证据、Reviewer、排序、TargetCards、报告与前端回放均有代码和契约。",
        "数据覆盖已跨疾病、组织、细胞和时期，但 disease library 的参考靶点主要来自已知药物与遗传证据，适合回归，不等于发现能力。",
        "最明显的运行缺口是 archive 全部 completed_with_gaps，且组学、扰动、安全三维在 fake/unit 排名中为 0。",
        "下一步应把 Reviewer findings 转成可追踪的 resolved 状态，并让真实工具证据进入六维排名。",
        "发布验收仍需 Python 3.11、live API 和 Reviewer LoRA GPU profile。",
    ]:
        base.add_bullet(doc, text)

    end_break = doc.add_paragraph(); end_break.add_run().add_break(WD_BREAK.PAGE)
    new_elements = []
    collecting = False
    for element in list(body):
        if element is marker_el:
            collecting = True
        if collecting and element.tag != base.qn("w:sectPr"):
            new_elements.append(element)
    for element in new_elements:
        body.remove(element)
    for element in new_elements:
        appendix._p.addprevious(element)
    body.remove(marker_el) if marker_el.getparent() is body else None

    doc.core_properties.title = "Target 项目全景与完整评测报告"
    doc.core_properties.subject = "Architecture, data assets, workflow artifacts, alignment, biology snapshots, and benchmarks"
    doc.core_properties.comments = "Generated from committed project assets and stored fake/unit runs; no live biological claims added."
    doc.save(DOCX)


def main() -> int:
    if not BASE_DOCX.exists():
        bm = base.collect_metrics(); base.create_docx(bm, base.create_figures(bm))
    metrics = collect_project_metrics()
    figs = create_project_figures(metrics)
    build_docx(metrics, figs)
    print(DOCX)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
